import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_ROW_HEIGHT_RULE
from io import BytesIO
import re

# --- 設定頁面資訊 ---
st.set_page_config(
    page_title="生日賀卡標籤生成器",
    page_icon="🏷️",
    layout="centered"
)

# --- 輔助函式 ---

def load_excel_with_auto_header(file):
    """
    自動偵測 Excel 的標題列位置。
    解決第一列是標題名稱(如: 臺東縣...)而不是欄位名稱的問題。
    """
    # 1. 先讀取前 10 列來預覽
    try:
        df_temp = pd.read_excel(file, header=None, nrows=10, dtype=str)
    except Exception:
        # 如果讀取失敗，回傳 None 讓後面處理
        return None
    
    header_idx = -1
    
    # 2. 逐列檢查是否包含關鍵欄位
    for idx, row in df_temp.iterrows():
        # 將整列轉為字串搜尋
        row_values = [str(val).strip() for val in row.values]
        if '姓名' in row_values and '通訊地址' in row_values:
            header_idx = idx
            break
            
    # 3. 重設檔案指標回到開頭
    file.seek(0)
    
    # 4. 根據找到的索引重新讀取
    if header_idx != -1:
        return pd.read_excel(file, header=header_idx, dtype=str)
    else:
        # 找不到關鍵字，就嘗試用預設方式讀取
        return pd.read_excel(file, dtype=str)

def process_address(raw_address):
    """
    處理地址邏輯：
    1. 嘗試從地址中提取郵遞區號 (例如: (950)臺東縣... -> 950, 臺東縣...)
    2. 如果沒有，則使用關鍵字對照表 (花蓮邏輯)
    """
    if not isinstance(raw_address, str):
        return "   ", ""

    raw_address = raw_address.strip()
    
    # 嘗試偵測開頭是否為 (數字) 或 數字
    # Regex 抓取開頭的 3碼數字，可能包含括號
    match = re.match(r'^[\(（]?(\d{3})[\)）]?(.*)', raw_address)
    
    if match:
        zip_code = match.group(1)
        clean_addr = match.group(2).strip()
        return zip_code, clean_addr
    
    # 如果地址本身沒有郵遞區號，則使用舊的對照表邏輯 (備用)
    zip_map = {
        "花蓮市": "970", "新城鄉": "971", "秀林鄉": "972",
        "吉安鄉": "973", "壽豐鄉": "974", "鳳林鎮": "975",
        "光復鄉": "976", "豐濱鄉": "977", "瑞穗鄉": "978",
        "萬榮鄉": "979", "玉里鎮": "981", "卓溪鄉": "982",
        "富里鄉": "983",
        "臺東市": "950" # 簡單補一個台東市，避免全空
    }
    
    found_zip = "   "
    for town, code in zip_map.items():
        if town in raw_address:
            found_zip = code
            break
            
    return found_zip, raw_address

def set_font(run, size=12, bold=False):
    """設定字型為標楷體 (中文) 與 Times New Roman (西文)"""
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(size)
    run.font.bold = bold

def generate_word_doc(df):
    """生成 Word 文件的核心邏輯"""
    doc = Document()
    
    # 設定版面: A4 大小，邊界全為 0
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    # 建立表格 (2欄 x N列)
    total_items = len(df)
    rows_needed = (total_items + 1) // 2 
    table = doc.add_table(rows=rows_needed, cols=2)
    
    for index, row_data in df.iterrows():
        r = index // 2
        c = index % 2
        
        # 取得資料並轉為字串
        name = str(row_data.get('姓名', '')).strip()
        raw_address = str(row_data.get('通訊地址', '')).strip()
        
        if name == 'nan': name = ''
        if raw_address == 'nan': raw_address = ''
        
        # 處理郵遞區號與地址
        zip_code, clean_address = process_address(raw_address)
        
        # 地址拆分邏輯：拆分 縣市鄉鎮 / 詳細地址
        # 簡單邏輯：取前 6 個字當鄉鎮 (例如: 臺東縣臺東市)，剩下當詳細地址
        # 這樣可以避免 (950) 佔用字數導致換行錯誤
        if len(clean_address) > 6:
            township = clean_address[:6]
            detail_addr = clean_address[6:]
        else:
            township = clean_address
            detail_addr = ""

        cell = table.rows[r].cells[c]
        cell.width = Cm(10.5)
        
        # 固定列高
        table.rows[r].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        table.rows[r].height = Cm(2.97) 

        cell.vertical_alignment = 1 # 垂直置中
        cell._element.clear_content()
        
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(10)
        
        # 第一行：郵遞區號 + 鄉鎮
        run1 = p.add_run(f"{zip_code} {township}\n")
        set_font(run1)
        
        # 第二行：詳細地址
        run2 = p.add_run(f"{detail_addr}\n")
        set_font(run2)
        
        # 第三行：姓名 + 稱謂
        if name:
            run3 = p.add_run(f"{name} 先生/女士 收") 
            set_font(run3, size=14, bold=True)

    # 將檔案存入記憶體 Buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Streamlit UI 介面 ---

st.title("🏷️ 生日賀卡標籤生成器")
st.markdown("""
本工具可將 Excel 通訊錄轉換為 **3M 21320 (A4 2欄 x 10列)** 格式的 Word 標籤檔。
請上傳您的 Excel 檔案進行轉換。
""")

st.info("💡 **提示**：程式會自動搜尋包含 **「姓名」** 與 **「通訊地址」** 的標題列。")

# 1. 檔案上傳區
uploaded_file = st.file_uploader("上傳 Excel 檔案 (.xlsx)", type=['xlsx'])

if uploaded_file is not None:
    try:
        # 使用新的讀取函式 (自動偵測標題)
        df = load_excel_with_auto_header(uploaded_file)
        
        if df is None:
            st.error("無法讀取 Excel 檔案，請確認檔案格式。")
            st.stop()
        
        # 檢查欄位是否存在
        required_cols = {'姓名', '通訊地址'}
        # 清理欄位名稱 (移除空白)
        df.columns = [str(c).strip() for c in df.columns]
        
        if not required_cols.issubset(df.columns):
            st.error(f"錯誤：Excel 缺少必要欄位！\n偵測到的欄位：{list(df.columns)}\n請確認檔案中包含：{required_cols}")
            st.stop()
            
        # 顯示前 5 筆資料預覽
        st.subheader("📋 資料預覽 (前 5 筆)")
        st.dataframe(df[['姓名', '通訊地址']].head())
        st.write(f"共讀取到 {len(df)} 筆資料")

        # 2. 生成按鈕
        if st.button("🚀 開始生成標籤", type="primary"):
            with st.spinner('正在排版並生成 Word 檔...'):
                docx_buffer = generate_word_doc(df)
                
                st.success("✅ 標籤生成成功！請點擊下方按鈕下載。")
                
                # 3. 下載按鈕
                st.download_button(
                    label="📥 下載 Word 標籤檔 (.docx)",
                    data=docx_buffer,
                    file_name="生日賀卡標籤.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.warning("⚠️ **列印注意**：列印 PDF 或 Word 時，請務必選擇 **「實際大小 (Actual Size)」** 或縮放比例 **100%**，以免標籤位置跑掉。")

    except Exception as e:
        st.error(f"發生錯誤：{e}")

# Footer
st.markdown("---")
st.caption("Designed for automated label processing.")
